import os
import shutil
from dotenv import load_dotenv
from docx import Document
from langchain_core.documents import Document as LangChainDocument

# LangChain 核心组件
from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import DashScopeEmbeddings
from langchain_community.chat_models import ChatTongyi
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser

# 加载环境变量
load_dotenv()
api_key = os.getenv("DASHSCOPE_API_KEY")
if not api_key:
    raise ValueError("请设置 DASHSCOPE_API_KEY 环境变量或在 .env 文件中配置。")

# ========== 配置常量 ==========
DOCUMENTS_DIR = "./documents"          # 存放 Word 和 PDF 的目录
VECTOR_STORE_DIR = "./vector_store"    # 本地向量数据库目录
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
RETRIEVER_K = 4                        # 检索返回的文本块数量
# =============================

def load_docx_file(file_path: str):
    """使用 python-docx 加载单个 .docx 文件"""
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            full_text.append(" | ".join(row_data))
    return LangChainDocument(
        page_content="\n".join(full_text),
        metadata={"source": file_path}
    )

def load_doc_file(file_path: str):
    """使用 docx2txt 加载单个 .doc 文件"""
    import docx2txt
    text = docx2txt.process(file_path)
    return LangChainDocument(
        page_content=text,
        metadata={"source": file_path}
    )

# 支持的文件类型与对应的 Loader 函数
LOADER_MAPPING = {
    ".pdf": PyPDFLoader,
    ".docx": load_docx_file,
    ".doc": load_doc_file,
}

# 文本分块器（全局复用）
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""],
)

# 嵌入模型（全局复用）
embeddings = DashScopeEmbeddings(
    model="text-embedding-v2",
    dashscope_api_key=api_key,
)

# 大语言模型（全局复用）
llm = ChatTongyi(
    model="qwen-max",
    dashscope_api_key=api_key,
    temperature=0.3,
)

# 提示模板
prompt_template = ChatPromptTemplate.from_messages([
    ("system", """你是一个基于本地知识库的问答助手。请根据以下参考资料回答用户问题。
如果参考资料中没有相关信息，请直接说“根据现有资料无法回答该问题”。
回答时，请尽量引用资料中的原文，并注明来源文件名称。

【参考资料】
{context}

【用户问题】
{question}
"""),
    ("human", "{question}")
])


def load_documents(directory: str):
    """加载指定目录下所有 PDF 和 Word 文档，返回 LangChain Document 列表"""
    all_docs = []
    
    for ext, loader_func in LOADER_MAPPING.items():
        file_list = []
        for root, _, files in os.walk(directory):
            for file in files:
                if file.lower().endswith(ext):
                    file_list.append(os.path.join(root, file))
        
        for file_path in file_list:
            try:
                if ext == ".pdf":
                    loader = loader_func(file_path)
                    docs = loader.load()
                    all_docs.extend(docs)
                else:
                    doc = loader_func(file_path)
                    all_docs.append(doc)
            except Exception as e:
                print(f"加载文件 {os.path.basename(file_path)} 出错: {e}")
    
    return all_docs


def format_docs(docs):
    """将检索到的文档列表格式化为提示词中的文本"""
    formatted = []
    for i, doc in enumerate(docs):
        source = os.path.basename(doc.metadata.get("source", "未知文件"))
        formatted.append(f"[{i+1}] 来自文件《{source}》的片段：\n{doc.page_content}")
    return "\n\n".join(formatted)


def create_or_load_vector_store(rebuild=False):
    """创建或加载本地向量数据库，返回 retriever"""
    # 如果需要重建，先删除旧数据库
    if rebuild and os.path.exists(VECTOR_STORE_DIR):
        import time
        max_retries = 5
        
        # 先尝试重命名文件夹，这样即使有程序在使用，也不会阻止我们创建新的
        temp_dir = VECTOR_STORE_DIR + "_old"
        try:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            os.rename(VECTOR_STORE_DIR, temp_dir)
            print("已将旧数据库重命名，正在创建新数据库...")
        except OSError:
            pass  # 如果重命名失败，继续尝试直接删除
        
        # 尝试删除旧文件夹（包括重命名后的）
        for dir_to_remove in [VECTOR_STORE_DIR, temp_dir]:
            if not os.path.exists(dir_to_remove):
                continue
            for i in range(max_retries):
                try:
                    shutil.rmtree(dir_to_remove)
                    if dir_to_remove == VECTOR_STORE_DIR:
                        print("已清除旧的向量数据库。")
                    break
                except PermissionError:
                    if i < max_retries - 1:
                        print(f"文件被占用，正在重试... ({i+1}/{max_retries})")
                        time.sleep(1)
                    else:
                        if dir_to_remove == VECTOR_STORE_DIR:
                            print("警告：无法删除旧的向量数据库文件，可能被其他程序占用。")
                            print("请关闭其他程序后重试，或者手动删除 vector_store 文件夹。")

    if os.path.exists(VECTOR_STORE_DIR) and os.listdir(VECTOR_STORE_DIR):
        print("检测到已有向量数据库，直接加载...")
        vector_store = Chroma(
            persist_directory=VECTOR_STORE_DIR,
            embedding_function=embeddings,
        )
    else:
        if not os.path.exists(DOCUMENTS_DIR):
            os.makedirs(DOCUMENTS_DIR)
            print(f"已创建 {DOCUMENTS_DIR} 目录，请放入 Word 或 PDF 文件后重新运行。")
            exit(0)

        docs = load_documents(DOCUMENTS_DIR)
        if not docs:
            print("未找到任何文档，程序退出。")
            exit(0)

        chunks = text_splitter.split_documents(docs)

        vector_store = Chroma.from_documents(
            documents=chunks,
            embedding=embeddings,
            persist_directory=VECTOR_STORE_DIR,
        )
        vector_store.persist()

    retriever = vector_store.as_retriever(
        search_type="similarity",
        search_kwargs={"k": RETRIEVER_K},
    )
    return retriever


def build_rag_chain(retriever):
    """构建 RAG 链"""
    rag_chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt_template
        | llm
        | StrOutputParser()
    )
    return rag_chain


def main():
    """交互式问答循环"""
    print("=" * 50)
    print("本地知识库问答系统（基于通义千问 + Chroma）")
    print(f"文档目录: {DOCUMENTS_DIR}")
    print(f"向量数据库: {VECTOR_STORE_DIR}")
    print("输入 'exit' 退出程序")
    print("=" * 50)

    # 重建索引可通过命令行参数 --rebuild 触发（见文件末尾）
    import sys
    rebuild = "--rebuild" in sys.argv

    retriever = create_or_load_vector_store(rebuild=rebuild)
    rag_chain = build_rag_chain(retriever)

    while True:
        try:
            question = input("\n请输入您的问题: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\n再见！")
            break

        if question.lower() in ("exit", "quit"):
            print("再见！")
            break
        if not question:
            continue

        try:
            answer = rag_chain.invoke(question)
            print("\n【答案】")
            print(answer)
            print("\n" + "-" * 30)
        except Exception as e:
            print(f"调用大模型时出错: {e}")


if __name__ == "__main__":
    main()